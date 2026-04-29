from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re


OUT = "docs/OneVo-HR-Phase-1-End-to-End-Scope-DETAILED.docx"
ROOT = Path(__file__).resolve().parents[1]


MODULES = [
    {
        "name": "Infrastructure",
        "phase": "Phase 1",
        "tables": ["countries", "file_records", "tenants", "users"],
        "purpose": "Foundational tenant, user, country, and file registry data used by almost every other module.",
        "scope": [
            "Tenant records for each customer organisation, including subscription/status context and branding links.",
            "User identity records used by authentication and sessions.",
            "Country reference data for legal entities, leave policies, public holidays, and nationality.",
            "Central file registry for uploaded files, verification photos, documents, screenshots, receipts, and certificates.",
        ],
    },
    {
        "name": "Auth and Security",
        "phase": "Phase 1",
        "tables": [
            "audit_logs", "feature_access_grants", "gdpr_consent_records", "permissions",
            "role_permissions", "roles", "sessions", "user_permission_overrides", "user_roles",
        ],
        "purpose": "Login, MFA, SSO, sessions, audit logging, GDPR consent, and hybrid permission control.",
        "scope": [
            "JWT login with refresh/session tracking and revocation.",
            "TOTP MFA and tenant/role-level MFA enforcement.",
            "Google OAuth and Azure AD SSO configuration through shared platform settings.",
            "Custom tenant roles as permission templates, not hard-coded role logic.",
            "Per-user permission grants/revocations that override assigned roles.",
            "Feature/module access grants to roles or individual employees.",
            "Hierarchy-scoped data access through reporting-line filters.",
            "Audit log capture for write operations, login events, exports, and security-sensitive actions.",
            "GDPR consent records before monitoring or biometric processing is enabled.",
        ],
    },
    {
        "name": "Org Structure",
        "phase": "Phase 1",
        "tables": [
            "department_cost_centers", "departments", "job_families", "job_levels",
            "job_titles", "legal_entities", "office_locations", "team_members", "teams",
        ],
        "purpose": "Company hierarchy, legal entities, departments, jobs, teams, offices, and cost centers.",
        "scope": [
            "Multiple legal entities per tenant with country and registration context.",
            "Parent-child department hierarchy with department heads.",
            "Job families, levels, titles, and salary band context.",
            "Teams separate from department hierarchy for operational grouping.",
            "Cost centers linked to departments for finance/reporting context.",
            "Office locations used by presence, public holidays, verification, and policy scoping.",
        ],
    },
    {
        "name": "Core HR",
        "phase": "Phase 1",
        "tables": [
            "employee_addresses", "employee_bank_details", "employee_custom_fields", "employee_dependents",
            "employee_emergency_contacts", "employee_lifecycle_events", "employee_qualifications",
            "employee_salary_history", "employee_work_history", "employees", "offboarding_records",
            "onboarding_tasks", "onboarding_templates",
        ],
        "purpose": "Employee system of record from hiring through offboarding.",
        "scope": [
            "Employee profiles with personal, work, reporting, job, department, office, and status fields.",
            "Addresses, bank details, emergency contacts, dependents, custom fields, work history, qualifications.",
            "Onboarding templates and task generation by department.",
            "Offboarding records and checklists.",
            "Lifecycle events for hire, transfer, promotion, salary change, and termination.",
            "Salary history with effective date, approval, and reason.",
        ],
    },
    {
        "name": "Skills Core",
        "phase": "Phase 1 subset",
        "tables": ["employee_skills", "job_skill_requirements", "skill_categories", "skill_validation_requests", "skills"],
        "purpose": "Skill taxonomy, employee skill profiles, job requirements, and validation requests.",
        "scope": [
            "Tenant skill categories and skill definitions.",
            "Job-family skill requirements.",
            "Employee-declared and manager/validator-confirmed skills.",
            "Validation request workflow for employee skill claims.",
        ],
    },
    {
        "name": "Leave",
        "phase": "Phase 1",
        "tables": ["leave_balances_audit", "leave_entitlements", "leave_policies", "leave_requests", "leave_types"],
        "purpose": "Configurable leave policies, entitlements, requests, approvals, cancellations, and balance audit.",
        "scope": [
            "Leave types per tenant.",
            "Leave policies scoped by country and job level.",
            "Employee leave entitlement assignment.",
            "Request submission, approval/rejection, cancellation, supporting document linkage.",
            "Leave balance audit history.",
        ],
    },
    {
        "name": "Calendar",
        "phase": "Phase 1",
        "tables": ["calendar_events"],
        "purpose": "Shared calendar events and conflict detection support.",
        "scope": [
            "Company/team/employee calendar events.",
            "Leave-driven calendar entries.",
            "Conflict checks used by leave and scheduling workflows.",
        ],
    },
    {
        "name": "Configuration",
        "phase": "Phase 1",
        "tables": [
            "app_allowlist_audit", "app_allowlists", "employee_monitoring_overrides",
            "integration_connections", "monitoring_feature_toggles", "tenant_settings",
        ],
        "purpose": "Tenant settings, integrations, monitoring controls, app allowlists, and employee overrides.",
        "scope": [
            "Tenant-level settings and operational defaults.",
            "Monitoring feature toggles by tenant.",
            "Employee-specific monitoring overrides.",
            "Application allowlists and audit of allowlist changes.",
            "Integration connection records with encrypted credentials.",
        ],
    },
    {
        "name": "Agent Gateway",
        "phase": "Phase 1",
        "tables": ["agent_commands", "agent_health_logs", "agent_policies", "registered_agents"],
        "purpose": "Only backend entry point for desktop agent registration, policy, heartbeat, ingestion, and commands.",
        "scope": [
            "Device registration and device JWT issuance.",
            "Employee-to-device linking at tray-app login.",
            "Policy distribution to each agent.",
            "Heartbeat and health reporting every 60 seconds.",
            "High-throughput ingest endpoint for activity batches.",
            "SignalR command channel for capture, start/stop, pause/resume, and policy refresh.",
            "Fallback command polling when SignalR is disconnected.",
        ],
    },
    {
        "name": "Activity Monitoring",
        "phase": "Phase 1",
        "tables": [
            "activity_daily_summary", "activity_raw_buffer", "activity_snapshots", "application_categories",
            "application_usage", "browser_activity", "device_tracking", "meeting_sessions", "screenshots",
        ],
        "purpose": "Stores and aggregates desktop activity data from the monitoring agent.",
        "scope": [
            "Raw agent batch buffering and asynchronous processing.",
            "Activity snapshots with keyboard/mouse counts only; no keystroke content.",
            "Application usage and productivity/allowlist classification.",
            "Browser activity where enabled.",
            "Meeting detection by process matching in Phase 1.",
            "Optional screenshots as restricted blob data with metadata only in the database.",
            "Daily summaries for reporting, productivity analytics, and exception rules.",
            "Presence-window validation so data outside active sessions or during breaks is discarded.",
        ],
    },
    {
        "name": "Discrepancy Engine",
        "phase": "Phase 1",
        "tables": ["discrepancy_events", "wms_daily_time_logs"],
        "purpose": "Compares OneVo activity/presence data against WMS time logs and records mismatches.",
        "scope": [
            "Stores WMS daily time logs per employee.",
            "Records discrepancy events for review and analytics.",
            "Consumes aggregated activity and WMS bridge data.",
        ],
    },
    {
        "name": "Workforce Presence",
        "phase": "Phase 1",
        "tables": [
            "attendance_records", "break_records", "device_sessions", "employee_schedules",
            "overtime_records", "presence_sessions", "public_holidays", "roster_entries",
            "roster_periods", "shift_assignments", "shifts", "work_schedules",
        ],
        "purpose": "Attendance, presence sessions, shifts, schedules, breaks, overtime, rosters, and device sessions.",
        "scope": [
            "Live presence sessions from desktop agent and biometric events.",
            "Attendance records, corrections, and approvals.",
            "Break tracking and pause/resume integration with the agent.",
            "Shifts, work schedules, assignments, rosters, and public holidays.",
            "Overtime records and approvals.",
            "Device-session records tied to registered agents.",
        ],
    },
    {
        "name": "Exception Engine",
        "phase": "Phase 1",
        "tables": ["alert_acknowledgements", "escalation_chains", "exception_alerts", "exception_rules", "exception_schedules"],
        "purpose": "Rule-based exception detection, alerting, acknowledgement, and escalation chains.",
        "scope": [
            "Configurable exception rules and schedules.",
            "Alerts for monitoring, presence, verification, app allowlist, offline agent, and other anomalies.",
            "Acknowledgement and resolution records.",
            "Escalation chains for unacknowledged alerts.",
            "Remote capture request events sent through Agent Gateway.",
        ],
    },
    {
        "name": "Identity Verification",
        "phase": "Phase 1",
        "tables": [
            "biometric_audit_logs", "biometric_devices", "biometric_enrollments",
            "biometric_events", "verification_policies", "verification_records",
        ],
        "purpose": "Photo verification and biometric terminal management for identity checks.",
        "scope": [
            "Tenant verification policies for login, logout, interval, and on-demand checks.",
            "Verification records with confidence score, trigger, method, device, requested-by, and alert linkage.",
            "Temporary verification photos stored as restricted blob files.",
            "Biometric device registration and encrypted device API keys.",
            "Biometric enrollments with mandatory consent.",
            "Biometric events and audit logs for device health/tamper events.",
        ],
    },
    {
        "name": "Productivity Analytics",
        "phase": "Phase 1",
        "tables": [
            "daily_employee_report", "monthly_employee_report", "weekly_employee_report",
            "wms_productivity_snapshots", "workforce_snapshot",
        ],
        "purpose": "Employee/team productivity rollups, trend reporting, and WMS productivity snapshots.",
        "scope": [
            "Daily, weekly, and monthly employee productivity summaries.",
            "Workforce snapshots for team/company views.",
            "WMS productivity snapshots from bridge data.",
            "Analytics export support where permissions allow it.",
        ],
    },
    {
        "name": "Shared Platform",
        "phase": "Phase 1",
        "tables": [
            "api_keys", "approval_actions", "compliance_exports", "escalation_rules", "feature_flags",
            "hardware_terminals", "legal_holds", "notification_channels", "notification_templates",
            "payment_methods", "plan_features", "rate_limit_rules", "refresh_tokens",
            "retention_policies", "scheduled_tasks", "signalr_connections", "sso_providers",
            "subscription_invoices", "subscription_plans", "system_settings", "tenant_branding",
            "tenant_feature_flags", "tenant_subscriptions", "user_preferences", "webhook_deliveries",
            "webhook_endpoints", "workflow_definitions", "workflow_instances",
            "workflow_step_instances", "workflow_steps", "bridge_api_keys", "wms_role_mappings",
            "wms_tenant_links",
        ],
        "purpose": "Cross-cutting platform functions: subscriptions, feature flags, workflows, notifications, webhooks, WMS bridge links, SSO, retention, and tenant branding.",
        "scope": [
            "Subscription plans, tenant subscriptions, invoices, payment methods, and plan features.",
            "Tenant branding and tenant feature flags.",
            "System settings, rate limits, scheduled tasks, and user preferences.",
            "Notification templates and channels used by the Notifications module.",
            "Workflow definitions, workflow instances, steps, and approval actions.",
            "Compliance exports, legal holds, and retention policies.",
            "SSO providers with encrypted credentials.",
            "Webhook endpoints and delivery tracking.",
            "WMS tenant links, bridge API keys, and role mappings.",
        ],
    },
    {
        "name": "Notifications",
        "phase": "Phase 1",
        "tables": ["No own tables; uses notification_templates and notification_channels in Shared Platform"],
        "purpose": "In-app, email, webhook, and real-time notification delivery across platform events.",
        "scope": [
            "Notification dispatch for leave, exceptions, verification failures, agent events, workflow actions, and system notices.",
            "Template/channel configuration stored under Shared Platform.",
            "SignalR live updates where real-time UI is required.",
        ],
    },
    {
        "name": "Developer Platform",
        "phase": "Phase 1 internal module",
        "tables": [
            "dev_platform_accounts", "dev_platform_sessions", "agent_version_releases",
            "agent_deployment_rings", "agent_deployment_ring_assignments",
        ],
        "purpose": "Internal OneVo operator console for tenant operations, flags, audit, system config, and agent version rollout.",
        "scope": [
            "Separate console.onevo.io Next.js app for OneVo internal team only.",
            "Google OAuth only for approved @onevo.io accounts.",
            "Separate admin JWT issuer: onevo-platform-admin.",
            "Separate backend host/namespace through ONEVO.Admin.Api and /admin/v1/* endpoints.",
            "Tenant console, provisioning, subscriptions, impersonation, and tenant troubleshooting.",
            "Feature flag manager for global flags and tenant overrides.",
            "Agent version manager with release catalog and deployment rings.",
            "Audit console for cross-tenant audit log review by platform operators.",
            "System config editor for global defaults and tenant-specific overrides.",
            "App catalog manager for observed applications and global productivity classifications.",
        ],
    },
]


PERMISSIONS = [
    "activity:read", "activity:read:self", "admin:access", "admin:agents", "admin:audit", "admin:compliance",
    "admin:devices", "admin:read", "admin:roles", "admin:users", "admin:write", "agent:command",
    "agent:manage", "agent:read", "agent:register", "agent:view-health", "agents:manage",
    "alerts:manage", "analytics:export", "analytics:read", "analytics:view", "analytics:write",
    "approvals:read", "attendance:approve", "attendance:read", "attendance:read-own", "attendance:read-team",
    "attendance:write", "audit:read", "billing:manage", "billing:read", "branding:manage",
    "bridges:read", "bridges:write", "calendar:read", "calendar:write", "chat:read", "chat:write",
    "compliance:manage", "departments:read", "deployment:write", "devices:manage", "docs:read",
    "documents:manage", "documents:read", "documents:write", "employee:create", "employees:bulk-update",
    "employees:create", "employees:delete", "employees:read", "employees:read-own", "employees:read-team",
    "employees:update", "employees:write", "exceptions:acknowledge", "exceptions:manage",
    "exceptions:read", "exceptions:resolve", "exceptions:view", "expense:admin", "expense:approve",
    "expense:create", "expense:manage", "expense:read", "goals:read", "goals:write",
    "grievance:manage", "grievance:read", "grievance:write", "hr:read", "inbox:read",
    "integrations:manage", "leave:approve", "leave:create", "leave:manage", "leave:none",
    "leave:read", "leave:read-own", "leave:read-team", "leave:write", "monitoring:configure",
    "monitoring:update-settings", "monitoring:view-settings", "notifications:configure",
    "notifications:manage", "notifications:read", "org:manage", "org:read", "org:write",
    "overtime:read", "payroll:approve", "payroll:manage", "payroll:read", "payroll:run",
    "payroll:view", "payroll:view-salary", "payroll:write", "people:read", "performance:manage",
    "performance:read", "performance:read-team", "performance:write", "planning:read",
    "planning:write", "platform:admin", "project:create", "projects:read", "projects:write",
    "reports:create", "reports:manage", "reports:read", "roles:manage", "roles:read",
    "schedule:read", "settings:admin", "settings:alerts", "settings:billing", "settings:branding",
    "settings:integrations", "settings:manage", "settings:notifications", "settings:read",
    "settings:system", "settings:write", "skills:manage", "skills:read", "skills:validate",
    "skills:write", "skills:write-team", "sprint:manage", "task:assign", "tasks:read",
    "tasks:write", "teams:read", "time:read", "time:write", "users:manage", "users:read",
    "verification:configure", "verification:read", "verification:review", "verification:view",
    "wms:chat", "wms:okr", "wms:projects", "workforce:approve-overtime",
    "workforce:correct-attendance", "workforce:dashboard", "workforce:manage",
    "workforce:manage-biometric", "workforce:read", "workforce:view",
]


SIMPLE_WORKFLOWS = [
    {
        "title": "Tenant / Company Setup",
        "user": [
            "Create a new company account as a tenant.",
            "Enter company name, legal details, country, subscription plan, and status.",
            "Add tenant branding such as logo, colors, and display name.",
            "Enable or disable product modules based on the purchased plan.",
            "Configure retention, rate limits, SSO providers, and notification channels where required.",
        ],
        "system": [
            "Create the tenant record.",
            "Attach subscription plan and feature flags.",
            "Create tenant branding and settings records.",
            "Prepare default system settings, notification templates, and workflow defaults.",
            "Keep all tenant data separated by tenant_id.",
        ],
        "edge": [
            ["Tenant already exists", "Block duplicate setup or require operator review."],
            ["Subscription plan missing", "Keep tenant in provisioning status until plan is assigned."],
            ["Branding not provided", "Use default OneVo branding until tenant branding is configured."],
        ],
    },
    {
        "title": "User Invitation and First Login",
        "user": [
            "Invite a user by email.",
            "Assign one or more roles to the user.",
            "Send invitation or password setup email.",
            "User logs in with email/password or SSO.",
            "User completes MFA setup if required.",
        ],
        "system": [
            "Create user identity record.",
            "Create user_roles records for assigned roles.",
            "Resolve effective permissions from roles, feature grants, and overrides.",
            "Create a session and issue tenant JWT.",
            "Refresh permissions when the access token is refreshed.",
            "Log login and security events in audit logs.",
        ],
        "edge": [
            ["MFA required but not configured", "Force MFA setup before access."],
            ["User disabled", "Block login."],
            ["Tenant feature disabled", "Hide UI and reject backend access for that feature."],
        ],
    },
    {
        "title": "Role, Permission, and Individual Override Setup",
        "user": [
            "Create a custom role such as HR Manager, Team Lead, Payroll Officer, or Viewer.",
            "Select the exact permissions for that role.",
            "Assign the role to one or more employees.",
            "Grant or revoke an individual permission for a specific employee.",
            "Grant feature/module access to a role or individual employee.",
        ],
        "system": [
            "Store the role in roles.",
            "Store role permissions in role_permissions.",
            "Store assigned roles in user_roles.",
            "Store employee-specific grants/revokes in user_permission_overrides.",
            "Store module-level grants in feature_access_grants.",
            "Calculate effective permissions as role permissions plus grants minus revokes.",
            "Apply hierarchy scope when returning employee or workforce data.",
        ],
        "edge": [
            ["Permission revoked at employee level", "Employee-level revoke wins over role permission."],
            ["Feature not granted", "Permission is ineffective even if role contains it."],
            ["Manager tries to access outside hierarchy", "Backend returns no data or forbidden response."],
        ],
    },
    {
        "title": "Organisation Structure Setup",
        "user": [
            "Create legal entities for the company.",
            "Create office locations under legal entities.",
            "Create departments and subdepartments.",
            "Assign department heads.",
            "Create job families, job levels, and job titles.",
            "Create teams and assign team leads/members.",
            "Create cost centers and link them to departments.",
        ],
        "system": [
            "Store legal entities and link them to countries.",
            "Store departments using parent-child relationships.",
            "Store job family/level/title records.",
            "Store teams and team_members.",
            "Store office locations and cost centers.",
            "Use this structure for employee profile, leave policy, reporting scope, team assignment, and workflow approvals.",
        ],
        "edge": [
            ["Department has children", "Keep hierarchy intact; do not delete without reassignment."],
            ["Employee outside manager scope", "Do not allow assignment unless user has required scope or bypass."],
            ["Legal entity missing", "Block employee/legal setup that depends on it."],
        ],
    },
    {
        "title": "Employee Profile Creation",
        "user": [
            "Create the employee profile after hiring decision.",
            "Enter personal details, work email, employee number, phone, DOB, gender, and nationality.",
            "Assign department, job title, manager, legal entity, office, employment type, work mode, hire date, and probation date.",
            "Upload employee photo/avatar if available.",
            "Add addresses, emergency contacts, dependents, bank details, qualifications, and work history.",
        ],
        "system": [
            "Create employees record.",
            "Link employee to user account where applicable.",
            "Store related records in addresses, emergency contacts, dependents, bank details, qualifications, and work history tables.",
            "Encrypt bank account number before storage.",
            "Create lifecycle event for hire.",
            "Make employee available to onboarding, leave, workforce presence, skills, and monitoring modules.",
        ],
        "edge": [
            ["Duplicate employee number", "Reject because employee number must be unique per tenant."],
            ["Missing required job/department", "Block profile creation or keep draft until complete."],
            ["Sensitive data entered", "Store under PII controls and avoid logging sensitive values."],
        ],
    },
    {
        "title": "Employee Onboarding",
        "user": [
            "Create the new employee profile first.",
            "Start onboarding for that employee.",
            "Choose the department-specific onboarding template if one exists.",
            "Use the global onboarding template if the department has no template.",
            "Review generated tasks for documentation, equipment, training, access, and orientation.",
            "Assign tasks to HR, IT, manager, or another responsible person.",
            "Track task due dates and completion status.",
            "Let assigned users complete only their own tasks.",
            "Close onboarding when every required task is complete.",
        ],
        "system": [
            "Receive employee id from the onboarding request.",
            "Check employees:write permission.",
            "Load employee record.",
            "Find department onboarding template.",
            "If department template is missing, find global template.",
            "If no template exists, return configured-template error.",
            "Create onboarding_tasks from template.tasks_json.",
            "Set each due date from hire date plus offset_days.",
            "Publish EmployeeOnboardingStarted.",
            "Notify task assignees.",
            "When a task is completed, verify caller is assigned_to.",
            "Set task status to completed and completed_at timestamp.",
            "When all tasks are complete, mark onboarding complete.",
        ],
        "edge": [
            ["No department template", "Use global template."],
            ["No global template", "Return 400: no onboarding template configured."],
            ["Task already completed", "Return 400."],
            ["Caller is not assigned_to", "Return 403."],
            ["All tasks completed", "Mark onboarding complete."],
        ],
    },
    {
        "title": "Employee Promotion",
        "user": [
            "Open employee profile.",
            "Choose promote action.",
            "Select new job title and effective date.",
            "Enter new salary if salary changes.",
            "Enter reason for promotion.",
            "Submit for approval or apply based on permission/workflow.",
        ],
        "system": [
            "Validate employee exists and is active.",
            "Update job_title_id when promotion is effective.",
            "Create salary history record if salary changes.",
            "Create employee_lifecycle_events record with event_type promoted.",
            "Notify employee, manager, and HR where configured.",
            "Expose promotion in employee lifecycle timeline.",
        ],
        "edge": [
            ["Backdated promotion", "Allow if HR correction policy permits; created_at still records actual insertion time."],
            ["Salary not changed", "Create promotion lifecycle event without salary history change."],
            ["Invalid job title", "Reject request."],
        ],
    },
    {
        "title": "Employee Transfer",
        "user": [
            "Open employee profile.",
            "Choose transfer action.",
            "Select new department, team, manager, office, or legal entity as needed.",
            "Set effective date and reason.",
            "Submit transfer.",
        ],
        "system": [
            "Validate employee exists and is active.",
            "Validate target department/team/manager/legal entity are accessible and valid.",
            "Update employee department/team/manager/location fields.",
            "Create employee_lifecycle_events record with event_type transferred.",
            "Notify old and new managers where configured.",
            "Downstream modules can update shift, leave policy, monitoring office context, and reporting scope.",
        ],
        "edge": [
            ["Target manager outside allowed scope", "Reject unless authorized."],
            ["Transfer affects leave/shift", "Downstream modules recalculate or require admin review."],
            ["Same department selected", "Reject or treat as no-op depending on UI rule."],
        ],
    },
    {
        "title": "Employee Offboarding",
        "user": [
            "Start offboarding from the employee profile.",
            "Select reason: resignation, termination, retirement, contract_end, or configured reason.",
            "Set last working date.",
            "Record knowledge risk level.",
            "Create or review checklist items for IT, HR, Finance, Manager, and Admin.",
            "Manager completes knowledge transfer and task reassignment.",
            "For high/critical risk, complete knowledge transfer or request/approve bypass.",
            "If bypass is approved, enter reason and optional penalty amount.",
            "Record exit interview notes.",
            "Review penalties_json for loans, notice-period penalties, asset recovery, and knowledge-transfer bypass penalties.",
            "Complete offboarding after all required steps are done.",
        ],
        "system": [
            "Check employees:write permission.",
            "Load employee and validate active/offboarding eligibility.",
            "Create offboarding_records with reason, last_working_date, knowledge_risk_level, status, and penalties_json.",
            "Create offboarding checklist.",
            "Make knowledge transfer mandatory for high/critical risk.",
            "Trigger additional handover workflow for critical risk.",
            "If bypass is requested, require HR/Admin or configured approver.",
            "Require bypass reason.",
            "Resolve penalty amount from manual input, tenant default, or zero audit-only.",
            "Append knowledge_transfer_bypass item to penalties_json.",
            "Before completion, validate knowledge transfer is complete or approved bypass exists.",
            "Deactivate user account.",
            "Revoke sessions and refresh tokens.",
            "Trigger agent revocation.",
            "Set employee status to terminated/resigned and record termination date.",
            "Publish termination/offboarding events.",
        ],
        "edge": [
            ["Employee not active", "Offboarding start is rejected."],
            ["Already offboarding", "Block duplicate offboarding."],
            ["High/critical handover incomplete", "Block completion until completed or bypassed."],
            ["Bypass without reason", "Block request."],
            ["Penalty amount without currency", "Block request."],
            ["Agent installed", "Revoke agent access."],
        ],
    },
    {
        "title": "Leave Type and Leave Policy Setup",
        "user": [
            "Create leave types such as annual, sick, maternity, unpaid, or custom types.",
            "Create leave policies by country, legal entity, job level, or tenant rule.",
            "Set entitlement amount, accrual rule, carry-forward rule, approval rule, and document requirement.",
            "Assign entitlements to employees.",
        ],
        "system": [
            "Store leave_types.",
            "Store leave_policies with country/job-level links where applicable.",
            "Create leave_entitlements for employees.",
            "Track changes in leave_balances_audit.",
            "Use policies when leave requests are submitted.",
        ],
        "edge": [
            ["Policy missing", "Block entitlement assignment or require manual entitlement."],
            ["Country/job level mismatch", "Use matching policy or fallback based on configured rules."],
            ["Carry-forward exceeded", "Cap balance according to policy."],
        ],
    },
    {
        "title": "Leave Request and Approval",
        "user": [
            "Employee selects leave type and date range.",
            "Employee adds reason and supporting document if required.",
            "Employee submits request.",
            "Manager or approver reviews request.",
            "Approver approves or rejects.",
            "Employee checks updated leave balance and calendar.",
        ],
        "system": [
            "Validate employee and leave type.",
            "Check available entitlement/balance.",
            "Check date conflicts and calendar context.",
            "Create leave_requests record.",
            "Notify approver.",
            "On approval, update entitlement/balance and audit trail.",
            "Create calendar event for approved leave.",
            "Notify employee.",
            "Pass absence context to workforce presence and payroll/final settlement where applicable.",
        ],
        "edge": [
            ["Insufficient balance", "Reject or require unpaid leave path."],
            ["Document required but missing", "Block submission."],
            ["Approver unavailable", "Use workflow escalation rules."],
            ["Request cancelled", "Reverse/adjust balance and calendar entry."],
        ],
    },
    {
        "title": "Calendar Event Creation",
        "user": [
            "Create company, department, team, or individual calendar event.",
            "Set title, date/time, audience, location, and description.",
            "Add participants where needed.",
            "Save event.",
        ],
        "system": [
            "Validate calendar:write permission.",
            "Validate audience and participant scope.",
            "Create calendar_events record.",
            "Run conflict checks where applicable.",
            "Notify participants where configured.",
            "Expose event to leave, attendance, and dashboard views.",
        ],
        "edge": [
            ["Participant outside scope", "Reject unless bypass/access exists."],
            ["Time conflict", "Warn or block based on event type."],
            ["Leave-driven event", "Created automatically from approved leave."],
        ],
    },
    {
        "title": "Monitoring Configuration",
        "user": [
            "Enable or disable activity monitoring by tenant.",
            "Enable or disable screenshot capture, identity verification, app tracking, device tracking, meeting detection, and browser/communication/document tracking where available.",
            "Set idle threshold, snapshot interval, heartbeat interval, and verification interval.",
            "Configure application allowlist and violation threshold.",
            "Create employee-level monitoring overrides.",
        ],
        "system": [
            "Store tenant defaults in monitoring_feature_toggles.",
            "Store employee overrides in employee_monitoring_overrides.",
            "Store app allowlists and audit changes.",
            "Merge tenant policy and employee override to produce effective agent policy.",
            "Push or refresh policy through Agent Gateway.",
            "Double-check policy on the server before processing incoming activity.",
        ],
        "edge": [
            ["Monitoring disabled", "Do not process incoming monitoring data for that employee/feature."],
            ["Employee override exists", "Override wins over tenant default."],
            ["Policy changed while agent online", "Send refresh policy command or agent fetches on next cycle."],
        ],
    },
    {
        "title": "Desktop Agent Registration and Login",
        "user": [
            "Install the Windows agent package on employee device.",
            "Employee opens tray app and logs in.",
            "Employee sees status and monitoring state in tray app.",
            "Employee can logout from tray app.",
        ],
        "system": [
            "Register device through Agent Gateway.",
            "Create registered_agents record.",
            "Issue device JWT separate from user JWT.",
            "On employee login, link employee_id to registered agent.",
            "Fetch current monitoring policy.",
            "Start enabled collectors.",
            "Send heartbeat every 60 seconds.",
            "Continue heartbeating after logout but stop employee data collection.",
        ],
        "edge": [
            ["Device offline", "Buffer locally and sync later."],
            ["Agent token invalid", "Reject requests until re-registered/logged in."],
            ["Employee offboarded", "Revoke agent registration/access."],
        ],
    },
    {
        "title": "Agent Data Ingestion",
        "user": [
            "No normal user action; agent sends data automatically while monitoring is active.",
            "Admins/managers view processed activity in dashboards if they have permission.",
        ],
        "system": [
            "Agent writes collected data to local SQLite buffer.",
            "Agent batches data and sends to POST /api/v1/agent/ingest.",
            "Agent Gateway validates device JWT.",
            "Gateway returns 202 Accepted quickly.",
            "Raw payload is stored in activity_raw_buffer.",
            "ProcessRawBufferJob parses data into activity_snapshots, application_usage, meeting_sessions, and device_tracking.",
            "Server validates presence window and break status.",
            "Data during breaks or outside active sessions is discarded.",
            "Daily aggregation creates activity_daily_summary.",
        ],
        "edge": [
            ["Network unavailable", "Agent stores data locally and retries."],
            ["Payload outside presence window", "Discard with warning log."],
            ["Raw buffer older than retention", "Purge after 48 hours."],
        ],
    },
    {
        "title": "Screenshot / On-Demand Capture",
        "user": [
            "Manager opens an exception alert.",
            "Manager requests screenshot or photo capture where permission allows.",
            "Employee receives notification before capture.",
            "Manager views result in alert detail after capture completes.",
        ],
        "system": [
            "Check agent:command permission.",
            "Create agent command with command_type capture_screenshot or capture_photo.",
            "Send command through SignalR agent hub.",
            "If SignalR is unavailable, agent picks command during polling.",
            "Agent shows employee notification and delay/window as required.",
            "Agent captures screenshot/photo and uploads file.",
            "Agent reports command completion.",
            "Identity Verification or Activity Monitoring stores metadata and links file_record.",
            "Notify manager result is available.",
        ],
        "edge": [
            ["Agent offline", "Command expires and manager must retry."],
            ["Capture rate limit exceeded", "Block request."],
            ["Employee notification required", "Capture must not proceed silently in Phase 1 flow."],
        ],
    },
    {
        "title": "AWS Rekognition Image Processing Plan",
        "user": [
            "Admin enables image processing policy where allowed.",
            "Manager/HR reviews verification result or exception-linked capture result.",
        ],
        "system": [
            "Receive image/photo from monitoring agent.",
            "Store image in tenant-scoped blob storage and file_records.",
            "Submit object reference to AWS Rekognition worker.",
            "Run configured Rekognition operation such as face comparison, face detection, or moderation/classification.",
            "Normalize result into confidence/status/failure reason.",
            "Store result against verification_records or alert context.",
            "Trigger exception/notification if verification fails or confidence is below threshold.",
            "Delete original image and metadata according to tenant retention policy.",
        ],
        "edge": [
            ["Consent/policy missing", "Do not submit image for processing."],
            ["Rekognition unavailable", "Mark processing pending/failed and retry according to job policy."],
            ["Legal hold active", "Do not delete retained image until hold is released."],
        ],
    },
    {
        "title": "Identity Verification",
        "user": [
            "Admin configures verification policy.",
            "Employee completes photo verification at login, logout, interval, or on-demand request.",
            "Manager/HR reviews failed verification alerts.",
            "Admin registers biometric devices where used.",
            "Admin enrolls employee fingerprint with consent where biometric hardware is used.",
        ],
        "system": [
            "Store verification_policies.",
            "Trigger photo capture based on policy.",
            "Store verification photo as restricted blob data.",
            "Create verification_records with method, trigger, status, confidence, device, and file reference.",
            "Publish verification completed or failed event.",
            "Create exception alert and notification on failed verification.",
            "For biometric devices, validate HMAC webhook and store biometric_events.",
        ],
        "edge": [
            ["Employee override disables verification", "Skip verification for that employee."],
            ["No consent for biometric enrollment", "Block biometric enrollment."],
            ["Photo expired", "Delete according to retention policy."],
        ],
    },
    {
        "title": "Workforce Presence and Attendance",
        "user": [
            "Employee clocks in/out or logs into active work session through agent/biometric flow.",
            "Employee starts and ends breaks.",
            "Manager views live presence dashboard.",
            "Manager or HR reviews attendance records.",
            "Manager/HR corrects attendance or approves overtime where permitted.",
        ],
        "system": [
            "Create presence_sessions and device_sessions.",
            "Create attendance_records for daily attendance.",
            "Create break_records for breaks.",
            "Pause agent data collection during breaks.",
            "Resume monitoring after break ends.",
            "Apply shifts, schedules, rosters, public holidays, and overtime rules.",
            "Send live updates through SignalR.",
            "Expose attendance corrections and overtime approval workflow.",
        ],
        "edge": [
            ["Break active", "Discard monitoring data during break."],
            ["Missing clock-out", "Flag attendance exception."],
            ["Overtime not approved", "Keep overtime pending until approval."],
        ],
    },
    {
        "title": "Application Allowlist Violation",
        "user": [
            "Admin configures allowed/non-allowed applications.",
            "Manager views application usage and violation alerts.",
            "Manager/HR acknowledges or resolves alert.",
        ],
        "system": [
            "During raw activity processing, compare app usage against resolved allowlist.",
            "Set is_allowed on application_usage.",
            "If non-allowed usage exceeds violation threshold, create/publish violation event.",
            "Exception Engine creates alert.",
            "Notifications alert manager or configured recipients.",
            "Manager acknowledgement is stored.",
        ],
        "edge": [
            ["Unknown app", "Surface in observed/app catalog for classification."],
            ["Employee override allows app", "Override is respected."],
            ["Threshold not reached", "Store usage but do not alert."],
        ],
    },
    {
        "title": "Exception Rule and Alert Handling",
        "user": [
            "Admin creates exception rules and escalation chains.",
            "Manager views exception dashboard.",
            "Manager opens alert detail.",
            "Manager acknowledges, dismisses, escalates, or resolves alert with notes.",
            "Manager may request on-demand capture if allowed.",
        ],
        "system": [
            "Store exception_rules and escalation_chains.",
            "Evaluate rules against activity, presence, verification, and agent health signals.",
            "Create exception_alerts.",
            "Notify assigned recipients.",
            "Record alert acknowledgements.",
            "Escalate if alert is not acknowledged in configured time.",
            "Link remote capture result to alert where requested.",
        ],
        "edge": [
            ["Rule inactive", "Do not evaluate it."],
            ["Alert not acknowledged", "Escalate based on chain."],
            ["Duplicate signal", "Avoid duplicate or handle idempotently based on rule design."],
        ],
    },
    {
        "title": "Productivity Analytics",
        "user": [
            "Manager views daily, weekly, or monthly employee/team reports.",
            "Manager compares active time, idle time, meeting time, application usage, and trends.",
            "Authorized user exports analytics when analytics:export is granted.",
        ],
        "system": [
            "Aggregate activity_daily_summary into daily_employee_report.",
            "Generate weekly and monthly employee reports.",
            "Generate workforce_snapshot for team/company overview.",
            "Include WMS productivity snapshots where WMS integration is enabled.",
            "Apply tenant, permission, and hierarchy scope to analytics queries.",
        ],
        "edge": [
            ["No activity data", "Show empty/insufficient data state."],
            ["Employee outside scope", "Do not include in report."],
            ["Export permission missing", "Allow view but block export."],
        ],
    },
    {
        "title": "Discrepancy Detection with WMS",
        "user": [
            "Manager or HR reviews discrepancy events.",
            "Manager compares WMS time logs with OneVo presence/activity data.",
            "Manager resolves or follows up on discrepancies.",
        ],
        "system": [
            "Receive/store WMS daily time logs.",
            "Compare WMS time with activity/presence summaries.",
            "Create discrepancy_events where mismatch rules are met.",
            "Expose discrepancy data to exception/analytics workflows.",
        ],
        "edge": [
            ["WMS not enabled", "Skip WMS discrepancy checks."],
            ["Missing WMS log", "Mark as missing input or skip based on rule."],
            ["Mismatch below threshold", "Do not create discrepancy event."],
        ],
    },
    {
        "title": "Notifications",
        "user": [
            "Admin configures notification channels and templates.",
            "Users receive in-app/email/webhook notifications depending on event and preference.",
            "Users open notification inbox or follow notification action.",
        ],
        "system": [
            "Use notification_templates and notification_channels from Shared Platform.",
            "Receive events from leave, onboarding, offboarding, exceptions, verification, workflows, and agent health.",
            "Render message from template.",
            "Deliver through configured channels.",
            "Track delivery status where webhook/email delivery records exist.",
            "Send SignalR updates for live UI where needed.",
        ],
        "edge": [
            ["Channel disabled", "Skip that channel."],
            ["Template missing", "Use fallback or log configuration issue."],
            ["Delivery fails", "Record delivery failure and retry where configured."],
        ],
    },
    {
        "title": "Developer Platform - Tenant Operations",
        "user": [
            "OneVo operator logs into console.onevo.io with approved Google SSO.",
            "Operator searches tenant.",
            "Operator views tenant status, subscription, flags, settings, and recent audit/activity.",
            "Operator provisions tenant or updates tenant settings where authorized.",
            "Operator may start impersonation/debug flow if policy allows.",
        ],
        "system": [
            "Authenticate platform user with Google OAuth.",
            "Issue platform-admin JWT, separate from tenant JWT.",
            "Reject platform JWT at customer /api/v1/* endpoints.",
            "Use /admin/v1/* endpoints only.",
            "Call existing application/module interfaces instead of bypassing business boundaries.",
            "Audit platform operator actions.",
        ],
        "edge": [
            ["Non @onevo.io account", "Block access."],
            ["Admin API disabled", "Console cannot perform operations."],
            ["Tenant JWT used on admin endpoint", "Reject token."],
        ],
    },
    {
        "title": "Developer Platform - Agent Version Rollout",
        "user": [
            "Operator creates a new desktop agent version release.",
            "Operator assigns rollout ring such as internal/canary, beta, or stable.",
            "Operator assigns tenants to deployment rings.",
            "Operator monitors rollout status.",
            "Operator pauses or rolls back release where supported.",
        ],
        "system": [
            "Store agent_version_releases.",
            "Store agent_deployment_rings.",
            "Store agent_deployment_ring_assignments.",
            "Agent checks version on startup/heartbeat.",
            "If newer version exists for assigned ring, agent downloads installer/update package.",
            "Agent applies update based on update policy.",
        ],
        "edge": [
            ["Tenant not assigned to ring", "Use default/stable ring."],
            ["Update fails", "Report agent health error."],
            ["Critical patch", "Operator can prioritize rollout according to platform policy."],
        ],
    },
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        shade(hdr[i], "1F4E5F")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def number(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def clean_md(text):
    text = text.replace("\u2014", "-").replace("\u2013", "-").replace("\u2192", "->")
    text = text.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    text = re.sub(r"\[\[([^\]|]+)\|([^\]]+)\]\]", r"\2", text)
    text = re.sub(r"\[\[([^\]]+)\]\]", lambda m: m.group(1).split("/")[-1], text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("Integration Events (cross-module - RabbitMQ)", "Cross-capability Events")
    text = text.replace("Integration Events (cross-module \u2014 RabbitMQ)", "Cross-capability Events")
    text = text.replace("Integration Events (cross-module â€” RabbitMQ)", "Cross-capability Events")
    return text


def add_code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        para.paragraph_format.space_after = Pt(0)


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    add_table(doc, rows[0], rows[1:], widths=None)


def add_markdown_file(doc, path, base_level=2):
    rel = path.relative_to(ROOT).as_posix()
    h(doc, rel, base_level)
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = clean_md(raw)
    lines = raw.splitlines()
    in_code = False
    code_lines = []
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, code_lines)
            code_lines = []

    for line in lines:
        stripped = line.strip()
        if any(term in stripped.lower() for term in [
            "modular monolith", "not microservices", "masstransit", "rabbitmq",
            "single applicationdbcontext", "all 176 tables live",
        ]):
            continue
        if stripped.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not stripped or stripped == "---":
            continue
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            level = min(base_level + max(stripped.count("#") - 1, 0), 4)
            h(doc, title, level)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith(">"):
            para = doc.add_paragraph(stripped.lstrip("> ").strip())
            for run in para.runs:
                run.italic = True
        else:
            p(doc, stripped)
    flush_table()
    flush_code()


def detail_files():
    dirs = [
        "modules/infrastructure", "modules/auth", "modules/org-structure", "modules/core-hr",
        "modules/skills", "modules/leave", "modules/calendar", "modules/configuration",
        "modules/agent-gateway", "modules/activity-monitoring", "modules/discrepancy-engine",
        "modules/workforce-presence", "modules/exception-engine", "modules/identity-verification",
        "modules/productivity-analytics", "modules/shared-platform", "modules/notifications",
        "developer-platform",
    ]
    files = []
    for d in dirs:
        root = ROOT / d
        if not root.exists():
            continue
        files.extend(sorted(root.rglob("overview.md")))
        files.extend(sorted(root.rglob("end-to-end-logic.md")))
    seen = set()
    ordered = []
    for f in files:
        if f not in seen:
            seen.add(f)
            ordered.append(f)
    return ordered


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("OneVo-HR Phase 1 End-to-End Scope | Client Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 95)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 95)
    styles["Heading 3"].font.name = "Aptos"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.color.rgb = RGBColor(45, 45, 45)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OneVo-HR Phase 1\nEnd-to-End System Scope")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 95)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Including Developer Platform, Monitoring Agent, and Planned AWS Rekognition Image Processing")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph()
    add_table(doc, ["Field", "Value"], [
        ["Prepared for", "Client"],
        ["Prepared from", "OneVo-HR workspace documentation and module specifications"],
        ["Document date", "2026-04-29"],
        ["Phase", "Phase 1 active delivery scope"],
        ["Classification", "Client Confidential"],
        ["Important boundary", "No undocumented product features were added. The AWS Rekognition section is explicitly marked as planned integration for monitoring-agent images."],
    ], widths=[1.8, 5.5])

    h(doc, "Executive Summary", 1)
    p(doc, "OneVo-HR is a multi-tenant, white-label SaaS platform that combines HR management with workforce intelligence. Phase 1 delivers the operational foundation, HR lifecycle capabilities, leave, organisation structure, skills core, workforce presence, desktop monitoring, identity verification, exception handling, productivity analytics, configuration, notifications, WMS integration support, and the internal Developer Platform used by OneVo operators.")
    p(doc, "The system uses a .NET 9 Clean Architecture backend with clear layer boundaries, CQRS-style use cases, a Next.js frontend, PostgreSQL persistence, and a Windows monitoring agent. The Developer Platform is separate from the customer frontend and uses a dedicated admin API namespace and a separate JWT issuer.")

    add_table(doc, ["Area", "Phase 1 scope"], [
        ["Customer product", "HR management plus workforce intelligence modules delivered through the main OneVo web application."],
        ["Monitoring agent", "Windows service plus tray app, device JWT, offline buffer, heartbeat, policy sync, activity ingestion, and command channel."],
        ["Developer platform", "Internal console.onevo.io operator app for tenants, flags, audit, config, app catalog, and agent version rollout."],
        ["Database scale", "133 Phase 1 tables across product modules, plus 5 Phase 1 Developer Platform tables."],
        ["Image processing plan", "Monitoring-agent images/photos will be routed to AWS Rekognition for planned face/image processing; database and retention remain tenant-scoped."],
    ], widths=[2.0, 5.3])

    h(doc, "Scope Rules Used In This Document", 1)
    bullets(doc, [
        "Only modules and behaviours present in the repository documentation were included as active Phase 1 scope.",
        "Phase 2 items are separated and marked as future scope where referenced.",
        "The planned AWS Rekognition flow is included because it was explicitly requested for images from the monitoring agent.",
        "Permission codes are listed explicitly from the available docs/code references instead of being summarized as a count.",
    ])

    h(doc, "Platform Architecture", 1)
    add_table(doc, ["Layer", "Technology / role"], [
        ["Frontend", "Next.js 14 App Router main customer application at app.onevo.io."],
        ["Backend", ".NET 9 Clean Architecture backend with CQRS-style application use cases."],
        ["Admin backend", "ONEVO.Admin.Api host for /admin/v1/* Developer Platform endpoints."],
        ["Database", "PostgreSQL 16 with tenant-scoped tables and row-level security design."],
        ["Real-time", "SignalR for live dashboards, alert updates, and agent command delivery."],
        ["Background jobs", "Hangfire for raw buffer processing, aggregation, retention, command expiry, and health checks."],
        ["Desktop agent", "Windows service + MAUI tray app + shared library, packaged for Windows Phase 1."],
        ["File storage", "S3-compatible blob storage in production; local disk for development."],
        ["WMS integration", "External Work Management System consumed via bridge contracts; WMS is not built by OneVo-HR Phase 1."],
    ], widths=[1.8, 5.5])

    p(doc, "Customer endpoints use tenant JWTs and /api/v1/* routes. Developer Platform endpoints use platform-admin JWTs and /admin/v1/* routes. A tenant JWT is rejected at admin endpoints, and a platform-admin JWT is rejected at tenant endpoints.")

    h(doc, "Product Configuration Matrix", 1)
    add_table(doc, ["Tier", "Included scope"], [
        ["HR Management", "Core foundation, HR pillar, shared platform, notifications, leave, org structure, skills core."],
        ["Work Management", "Core foundation and WMS bridge people sync only."],
        ["HR + Workforce Intel", "HR Management plus monitoring agent, activity monitoring, presence, verification, exceptions, analytics."],
        ["HR + Work Management", "HR Management plus all WMS bridge contracts."],
        ["Full Suite", "HR, Workforce Intelligence, WMS integration, and all shared platform capabilities."],
    ], widths=[2.0, 5.3])

    h(doc, "Backend Architecture and Delivery Model", 1)
    bullets(doc, [
        "The backend follows Clean Architecture as the primary architectural model.",
        "Business capability areas are organized with clear boundaries while following the same Clean Architecture layer rules.",
        "All business use cases follow command/query separation through CQRS-style handlers.",
        "Domain entities and domain events remain independent of infrastructure concerns.",
        "The infrastructure layer owns EF Core mappings, persistence, external integrations, and file storage adapters.",
        "The API layer exposes tenant-facing /api/v1/* endpoints and admin-facing /admin/v1/* endpoints through separate hosts/namespaces.",
        "Cross-capability interaction should happen through application-layer interfaces, domain events, or explicitly defined integration boundaries rather than bypassing layer rules.",
    ])
    add_table(doc, ["Backend concern", "Phase 1 approach"], [
        ["Tenant isolation", "tenant_id on tenant-scoped data, application tenant context, and PostgreSQL row-level-security design."],
        ["Feature boundaries", "Clean Architecture feature folders and module interfaces; admin API uses module interfaces rather than direct table ownership."],
        ["Validation", "Command validators and explicit Result-style outcomes for business failures."],
        ["Background work", "Hangfire jobs for ingestion processing, rollups, cleanup, retention, command expiry, and health checks."],
        ["Real-time updates", "SignalR connections for live dashboards, agent command delivery, alert updates, and presence status."],
        ["Storage", "PostgreSQL metadata plus S3-compatible blob storage for files/images."],
    ], widths=[2.0, 5.3])

    h(doc, "Frontend Scope", 1)
    bullets(doc, [
        "Main customer application uses Next.js App Router.",
        "Sidebar navigation is permission-gated; unavailable items are hidden rather than disabled.",
        "Workforce Intelligence UI requires both permissions and monitoring configuration.",
        "Frontend permission checks are UX gates only; backend still enforces authorization and data scope.",
        "Role-scoped UX distinguishes employee, manager, HR admin, and org owner experiences.",
        "The Developer Platform frontend is a separate Next.js application and does not share customer-app sessions.",
    ])

    h(doc, "Phase 1 Module-by-Module Scope", 1)
    for m in MODULES:
        h(doc, m["name"], 2)
        add_table(doc, ["Field", "Detail"], [
            ["Phase", m["phase"]],
            ["Purpose", m["purpose"]],
            ["Tables", ", ".join(m["tables"])],
        ], widths=[1.3, 6.0])
        bullets(doc, m["scope"])

    h(doc, "Tenant Access, RBAC, and Hierarchy Scoping", 1)
    p(doc, "OneVo does not rely on hard-coded tenant roles for access. Roles are configurable templates, and the actual authorization result is the effective permission set after role permissions, individual grants, individual revocations, feature/module grants, and hierarchy filters are applied.")
    number(doc, [
        "A Super Admin creates tenant-specific custom roles.",
        "Permissions are assigned to roles as templates.",
        "Users receive one or more roles.",
        "The system adds any employee-level grant overrides.",
        "The system removes any employee-level revoke overrides.",
        "The system filters permissions by feature_access_grants for role-level or employee-level feature/module access.",
        "The system applies hierarchy scope to data queries, such as own record only, subordinate records, department/subdepartment records, or all records for Super Admin.",
    ])
    add_table(doc, ["Scope type", "Meaning"], [
        ["Own only", "Employee sees only their own profile, leave, activity, or related records."],
        ["Team/subordinate scope", "Manager sees direct and indirect reports through reporting-line hierarchy."],
        ["Department scope", "Department head sees employees in their department and subdepartments."],
        ["All tenant scope", "Tenant Super Admin/Org Owner sees all tenant records subject to permission and feature grants."],
        ["Platform scope", "Developer Platform operators are not tenant users; they use platform-admin identity and admin endpoints."],
    ], widths=[2.0, 5.3])

    h(doc, "API Surface Summary", 1)
    p(doc, "The exact final controller inventory may expand as handlers are implemented, but the repository already defines the following concrete API contracts for critical Phase 1 areas.")
    add_table(doc, ["Area", "Route", "Permission/Auth", "Purpose"], [
        ["Auth/Roles", "GET /api/v1/roles", "roles:read", "List roles."],
        ["Auth/Roles", "POST /api/v1/roles", "roles:manage", "Create custom role."],
        ["Auth/Roles", "PUT /api/v1/roles/{id}", "roles:manage", "Update role."],
        ["Auth/Roles", "DELETE /api/v1/roles/{id}", "roles:manage", "Delete custom role."],
        ["Auth/Roles", "POST /api/v1/roles/{id}/permissions", "roles:manage", "Set role permissions."],
        ["Auth/User Roles", "POST /api/v1/users/{id}/roles", "roles:manage", "Assign role to employee/user."],
        ["Auth/User Roles", "DELETE /api/v1/users/{id}/roles/{roleId}", "roles:manage", "Remove role from user."],
        ["Auth/Permissions", "GET /api/v1/users/{id}/permissions", "roles:manage", "Get effective permissions."],
        ["Auth/Permissions", "POST /api/v1/users/{id}/permission-overrides", "roles:manage", "Grant/revoke individual permission."],
        ["Auth/Permissions", "DELETE /api/v1/users/{id}/permission-overrides/{permId}", "roles:manage", "Remove override."],
        ["Feature Access", "GET /api/v1/feature-access", "roles:manage", "List feature/module grants."],
        ["Feature Access", "POST /api/v1/feature-access", "roles:manage", "Grant feature access."],
        ["Feature Access", "DELETE /api/v1/feature-access/{id}", "roles:manage", "Revoke feature access."],
        ["Agent", "POST /api/v1/agent/register", "Tenant API key", "Register new device and receive device JWT."],
        ["Agent", "POST /api/v1/agent/heartbeat", "Device JWT", "Update heartbeat and health."],
        ["Agent", "GET /api/v1/agent/policy", "Device JWT", "Fetch monitoring policy."],
        ["Agent", "POST /api/v1/agent/ingest", "Device JWT", "Submit activity batch."],
        ["Agent", "POST /api/v1/agent/login", "Device JWT + employee credentials", "Link employee to device."],
        ["Agent", "POST /api/v1/agent/logout", "Device JWT", "Unlink employee from device."],
        ["Agent", "GET /api/v1/agent/commands", "Device JWT", "Fetch pending commands fallback."],
        ["Agent", "POST /api/v1/agent/commands/{id}/ack", "Device JWT", "Acknowledge command receipt."],
        ["Agent", "POST /api/v1/agent/commands/{id}/complete", "Device JWT", "Report command completion."],
        ["Manager Capture", "POST /api/v1/agents/{agentId}/capture-screenshot", "agent:command", "Request remote screenshot."],
        ["Manager Capture", "POST /api/v1/agents/{agentId}/capture-photo", "agent:command", "Request remote photo."],
        ["Activity", "GET /api/v1/activity/snapshots/{employeeId}", "workforce:view", "Activity snapshots for date range."],
        ["Activity", "GET /api/v1/activity/summary/{employeeId}", "workforce:view", "Daily summary."],
        ["Activity", "GET /api/v1/activity/apps/{employeeId}", "workforce:view", "Application usage breakdown."],
        ["Activity", "GET /api/v1/activity/meetings/{employeeId}", "workforce:view", "Meeting sessions."],
        ["Activity", "GET /api/v1/activity/screenshots/{employeeId}", "workforce:view", "Screenshot metadata."],
        ["Activity", "GET /api/v1/activity/screenshots/{id}/view", "workforce:view", "Redirect to screenshot blob URL."],
        ["Activity", "GET /api/v1/activity/categories", "monitoring:view-settings", "Read app categories."],
        ["Activity", "POST /api/v1/activity/categories", "monitoring:configure", "Create/update category."],
        ["Activity", "DELETE /api/v1/activity/categories/{id}", "monitoring:configure", "Delete category."],
        ["Identity", "GET /api/v1/verification/policy", "verification:view", "Get tenant verification policy."],
        ["Identity", "PUT /api/v1/verification/policy", "verification:configure", "Update policy."],
        ["Identity", "GET /api/v1/verification/records/{employeeId}", "verification:view", "Verification history."],
        ["Identity", "POST /api/v1/verification/verify", "Internal agent", "Submit verification photo."],
        ["Biometric", "GET /api/v1/biometric/devices", "verification:view", "List devices."],
        ["Biometric", "POST /api/v1/biometric/devices", "verification:configure", "Register device."],
        ["Biometric", "PUT /api/v1/biometric/devices/{id}", "verification:configure", "Update device."],
        ["Biometric", "POST /api/v1/biometric/enroll", "verification:configure", "Enroll fingerprint."],
        ["Biometric", "POST /api/v1/biometric/webhook", "HMAC-SHA256", "Receive terminal events."],
        ["Developer Platform", "/admin/v1/*", "PlatformAdmin policy", "Internal operator endpoints only."],
    ], widths=[1.3, 2.3, 1.7, 2.0])

    h(doc, "Monitoring Agent Deep Dive", 1)
    p(doc, "The WorkPulse Agent is the monitoring package installed on employee Windows devices in Phase 1. It is the source of desktop activity, application usage, device-session, screenshot/photo, heartbeat, and health telemetry.")
    add_table(doc, ["Component", "Purpose"], [
        ["ONEVO.Agent.Service", "Always-on Windows Service that runs collectors, buffers data locally, syncs data, sends heartbeat, stores device tokens, and detects tamper conditions."],
        ["ONEVO.Agent.TrayApp", "MAUI tray UI for employee login/logout, status display, break/pause interaction, notifications, and photo capture."],
        ["ONEVO.Agent.Shared", "Shared IPC messages, models, constants, and policy DTOs used by Service and TrayApp."],
        ["ONEVO.Agent.Installer", "MSIX packaging for Windows Phase 1 installation."],
    ], widths=[2.3, 5.0])
    h(doc, "Agent Data Collection", 2)
    add_table(doc, ["Collector / capability", "Phase 1 behaviour"], [
        ["Activity", "Keyboard and mouse event counts only; no keystroke content."],
        ["App tracking", "Foreground application detection through Windows APIs."],
        ["Idle detection", "Idle/active time detection through OS input timing."],
        ["Meeting detection", "Basic process-name matching for Teams, Zoom, Meet, or other meeting apps."],
        ["Device tracking", "Active/idle device cycles and device session context."],
        ["Document tracking", "Word, Excel, PowerPoint, and Google Docs time tracking where enabled."],
        ["Communication tracking", "Outlook/Slack active time and send-event counts where enabled."],
        ["Screenshot", "Restricted image capture; metadata in database, image file in blob storage. Phase 1 docs distinguish remote/on-demand command capture and optional screenshot metadata."],
        ["Photo verification", "Tray app camera capture for identity verification triggers."],
    ], widths=[2.2, 5.1])
    h(doc, "Agent-to-Server Flow", 2)
    number(doc, [
        "Agent installs and registers the device through POST /api/v1/agent/register using tenant-level registration context.",
        "Server creates a registered_agents record and issues a device JWT.",
        "Employee logs in through the tray app; the device is linked to the employee through POST /api/v1/agent/login.",
        "Agent fetches policy through GET /api/v1/agent/policy.",
        "Enabled collectors write records to local SQLite immediately.",
        "Every snapshot interval, default 150 seconds, the sync service batches data to POST /api/v1/agent/ingest.",
        "Agent sends POST /api/v1/agent/heartbeat every 60 seconds with health and buffer status.",
        "Backend returns 202 Accepted for ingest, then processes raw buffer asynchronously.",
        "SignalR /hubs/agent-commands pushes start, stop, pause, resume, refresh policy, screenshot, and photo commands.",
        "If SignalR drops, agent polls GET /api/v1/agent/commands during heartbeat cycles.",
    ])
    h(doc, "Agent Privacy and Control Rules", 2)
    bullets(doc, [
        "Device JWTs are separate from user JWTs and contain device/tenant context, not HR permissions.",
        "The agent cannot read HR data, employee profiles, payroll, leave, or other business records.",
        "Monitoring policy is computed from tenant toggles and employee overrides; the server validates again on ingest.",
        "Data collected during breaks is not processed into activity analytics.",
        "Window titles are hashed before storage.",
        "Screenshots and verification photos are restricted data and are stored as files, not database blobs.",
        "Employee notification is mandatory before on-demand screenshot/photo capture in the documented flow.",
        "On-demand capture requests are rate limited to prevent harassment.",
    ])

    h(doc, "Planned AWS Rekognition Image Processing", 1)
    p(doc, "This section is a planned integration for monitoring-agent images/photos. It is not described as already implemented. The purpose is to process images that originate from the monitoring agent, especially identity verification photos and on-demand capture results, using AWS Rekognition while keeping OneVo's tenant isolation, consent, retention, and audit rules intact.")
    add_table(doc, ["Step", "Planned processing path"], [
        ["1. Capture", "Agent captures a verification photo or on-demand screenshot/photo after policy and notification requirements are satisfied."],
        ["2. Upload", "Image is uploaded to tenant-scoped S3-compatible storage path and registered in file_records."],
        ["3. Event", "Agent Gateway reports command completion or image submission; Identity Verification creates/updates verification_records."],
        ["4. Rekognition request", "Backend image-processing worker submits the stored object reference to AWS Rekognition for face comparison, face detection, or image moderation/classification as configured."],
        ["5. Result mapping", "AWS response is normalized into match confidence, labels, or failure reasons and linked back to verification_records or exception alert context."],
        ["6. Alerting", "Low confidence, failed verification, or policy violations publish events to Exception Engine and Notifications."],
        ["7. Retention", "Original images and derived metadata follow tenant retention policies and legal hold rules."],
    ], widths=[1.6, 5.7])
    h(doc, "AWS Rekognition Controls", 2)
    bullets(doc, [
        "Use IAM role-based access with least privilege for Rekognition and object read access.",
        "Never send images unless tenant policy, employee consent, and monitoring override rules allow processing.",
        "Preserve tenant_id linkage in job metadata and audit logs.",
        "Do not log image content, raw faces, or extracted sensitive image details.",
        "Store only normalized result metadata required by OneVo workflows, such as confidence score and status.",
        "Respect default 30-day retention for verification photos/screenshots unless tenant policy says otherwise.",
        "Apply legal holds before deletion where compliance exports or investigations require preservation.",
    ])

    h(doc, "Developer Platform Module", 1)
    p(doc, "The Developer Platform is an internal control plane, not a customer developer portal. It runs as a separate Next.js app at console.onevo.io and uses Google OAuth for approved @onevo.io users only.")
    add_table(doc, ["Module", "Capabilities"], [
        ["Tenant Console", "Tenant search, provisioning, subscription status, feature/module enablement, tenant-level troubleshooting, and impersonation support."],
        ["Feature Flag Manager", "Global flags, tenant overrides, module enable/disable controls, and rollout visibility."],
        ["Desktop Agent Version Manager", "Agent release catalog, deployment rings, tenant ring assignments, and force/managed update operations."],
        ["Audit Console", "Cross-tenant audit search for platform operators, with filters by tenant, user, action, date, IP, and resource."],
        ["System Config", "Global defaults and tenant-specific overrides for values not exposed in the customer UI."],
        ["App Catalog Manager", "Observed application catalog, public/private visibility, global productivity suggestions, and bulk approval of uncatalogued apps."],
        ["Platform API Keys", "Phase 2 only; not part of Phase 1 active customer-facing developer API scope."],
    ], widths=[2.0, 5.3])
    h(doc, "Developer Platform Security Boundary", 2)
    bullets(doc, [
        "Separate frontend from the customer app.",
        "Separate domain: console.onevo.io.",
        "Separate auth provider and sessions: Google OAuth/NextAuth for internal platform users.",
        "Separate JWT issuer: onevo-platform-admin.",
        "Separate backend namespace: /admin/v1/*.",
        "Infrastructure-level VPN/IP allowlist gate is planned for network isolation.",
        "Admin API controllers call existing module interfaces; they do not bypass module boundaries or directly own customer module data.",
    ])

    h(doc, "WMS Integration Scope", 1)
    p(doc, "The Work Management System is consumed by OneVo but is not built as part of this OneVo-HR Phase 1 scope. OneVo owns bridge configuration, tenant links, role mappings, bridge API keys, and frontend consumption of WMS-backed views.")
    add_table(doc, ["WMS area", "OneVo responsibility"], [
        ["People sync", "Map OneVo employees/users into WMS people contracts for tiers that include Work Management."],
        ["Projects/tasks/sprints/OKRs/chat", "Consume WMS backend bridge contracts; do not implement the WMS backend itself."],
        ["Bridge API keys", "Store/manage tenant-scoped bridge keys through Shared Platform tables."],
        ["Role mappings", "Map OneVo roles to WMS roles through wms_role_mappings."],
        ["Tenant links", "Track WMS tenant link status and bridge configuration through wms_tenant_links."],
        ["Productivity snapshots", "Store WMS productivity snapshots for analytics where the tenant tier enables WMS integration."],
    ], widths=[2.0, 5.3])

    h(doc, "Security, Data Classification, and Retention", 1)
    add_table(doc, ["Data class", "Examples", "Control"], [
        ["Critical PII", "Bank account numbers, SSO client secrets, hardware terminal API keys, integration credentials.", "AES-256 encryption at rest for encrypted columns plus log scrubbing."],
        ["Restricted workforce data", "Screenshots, verification photos, application window title hash.", "Blob storage for image files, metadata only in database, access controlled by permission and tenant scope."],
        ["Confidential activity data", "Activity snapshots and raw agent buffer.", "Retention limits, partitioning, no content logging, tenant scoping."],
        ["Sensitive PII", "Emails, employee names, addresses, emergency contacts, dependents, session IPs.", "RLS/access control, targeted log scrubbing, audit trail where appropriate."],
    ], widths=[1.8, 2.5, 3.0])
    add_table(doc, ["Data category", "Retention"], [
        ["Active employee data", "While employed plus 7 years, then archive/delete according to policy."],
        ["Audit logs", "7 years."],
        ["Biometric events", "2 years."],
        ["Session data", "30 days after logout."],
        ["Notification records", "1 year."],
        ["Activity raw buffer", "48 hours."],
        ["Activity snapshots", "90 days."],
        ["Activity daily summaries", "2 years."],
        ["Screenshots", "Per tenant policy, default 30 days."],
        ["Verification photos", "Per tenant policy, default 30 days."],
        ["Agent health logs", "30 days."],
    ], widths=[2.5, 4.8])

    h(doc, "Operational Jobs and Monitoring", 1)
    add_table(doc, ["Job", "Schedule", "Purpose"], [
        ["ProcessRawBufferJob", "Every 2 minutes", "Parse raw agent buffer into snapshots, app usage, meetings, and tracking tables."],
        ["AggregateDailySummaryJob", "Every 30 minutes plus end-of-day", "Roll up snapshots into daily activity summaries."],
        ["PurgeRawBufferJob", "Daily", "Delete/drop raw buffer partitions older than 48 hours."],
        ["PurgeExpiredScreenshotsJob", "Daily", "Delete screenshot files and metadata past retention."],
        ["PurgeExpiredSnapshotsJob", "Monthly", "Drop snapshot partitions older than 90 days."],
        ["DetectOfflineAgentsJob", "Every 5 minutes", "Find agents with missing heartbeat and raise offline agent event/alert."],
        ["ExpirePendingCommandsJob", "Every 1 minute", "Expire stale pending agent commands."],
        ["CleanupRevokedAgentsJob", "Daily", "Clean old health logs for revoked agents."],
        ["CleanupCompletedCommandsJob", "Daily", "Clean completed/expired commands after retention window."],
        ["PurgeExpiredVerificationPhotosJob", "Daily", "Delete verification photos past retention."],
        ["CheckBiometricDeviceHealthJob", "Every 5 minutes", "Flag biometric devices without heartbeat."],
    ], widths=[2.1, 1.6, 3.6])

    h(doc, "End-to-End Business Flows", 1)
    flows = [
        ("Employee Onboarding", [
            "HR creates employee profile in Core HR.",
            "Auth creates/invites the user account.",
            "Role and permission templates are assigned.",
            "Leave entitlements are assigned from policy.",
            "Shift/schedule is assigned by Workforce Presence.",
            "Onboarding tasks are generated from the department template.",
            "Monitoring agent is installed and registered if workforce intelligence is enabled.",
            "Employee gives required consent before monitoring/biometric processing.",
            "Monitoring policy activates and welcome notification is sent.",
        ]),
        ("Leave Request and Approval", [
            "Employee submits leave request.",
            "Leave checks entitlement balance.",
            "Calendar checks conflict.",
            "Workforce Presence checks team/shift coverage context.",
            "Approver receives notification.",
            "Approver approves/rejects.",
            "Calendar and attendance/presence context are updated.",
            "Employee receives result notification.",
        ]),
        ("Monitoring Alert Escalation", [
            "Agent sends activity/presence data.",
            "Activity Monitoring processes raw buffer and daily summaries.",
            "Exception Engine evaluates configured rules.",
            "Alert is created and sent to manager dashboard.",
            "If unacknowledged, escalation chain advances.",
            "Manager/HR acknowledges or resolves with notes.",
        ]),
        ("On-Demand Capture", [
            "Manager opens exception alert and requests photo or screenshot.",
            "Exception Engine publishes remote capture request.",
            "Agent Gateway dispatches command over SignalR.",
            "Agent notifies employee before capture.",
            "Agent uploads image and reports command completion.",
            "Identity Verification records capture metadata and links alert/requester.",
            "Result becomes available to manager in alert detail.",
        ]),
        ("Employee Offboarding", [
            "HR initiates offboarding workflow.",
            "Core HR creates checklist and lifecycle record.",
            "Auth sessions are revoked after completion.",
            "Agent registration is revoked/deactivated.",
            "Leave balance closure is calculated.",
            "Employee status is set to terminated.",
            "Retention schedule is applied according to tenant policy.",
        ]),
    ]
    for title_text, steps in flows:
        h(doc, title_text, 2)
        number(doc, steps)

    h(doc, "Detailed Core HR Workflows", 1)
    h(doc, "Employee Onboarding - What HR Can Do", 2)
    bullets(doc, [
        "Create or select an employee profile as the onboarding subject.",
        "Start onboarding from POST /api/v1/employees/{id}/onboarding with employees:write permission.",
        "Use department-specific onboarding templates where configured.",
        "Fall back to a global onboarding template when the department does not have its own template.",
        "Generate onboarding tasks from template.tasks_json.",
        "Assign tasks to HR, IT, the manager, or another configured responsible user.",
        "Set due dates from the employee hire date plus each template task's offset_days.",
        "Track task categories including documentation, equipment, training, access, and orientation.",
        "Allow assigned users to complete their own assigned tasks.",
        "Automatically mark onboarding complete when all generated tasks are completed.",
        "Notify assignees through the EmployeeOnboardingStarted event and notification handling.",
    ])
    h(doc, "Employee Onboarding - System Sequence", 2)
    number(doc, [
        "EmployeeController.StartOnboarding receives the employee id.",
        "Authorization checks employees:write.",
        "OnboardingService loads the employee through the employee service.",
        "The service searches for an onboarding template matching department_id.",
        "If no department template exists, the service searches for a global template.",
        "If no usable template exists, the request fails with a configured-template error.",
        "For every template task, the service inserts an onboarding_tasks record.",
        "Each task stores task_name, category, assigned_to_id, due_date, and status.",
        "The service publishes EmployeeOnboardingStarted.",
        "Notifications alert the responsible task assignees.",
        "Task completion verifies the caller is the assigned person.",
        "Completed tasks receive status completed and completed_at timestamp.",
        "The service checks whether all tasks for the employee are complete.",
    ])
    add_table(doc, ["Onboarding error/edge case", "System handling"], [
        ["No department template", "Use global template."],
        ["No global template", "Return 400: no onboarding template configured."],
        ["Task already completed", "Return 400."],
        ["Caller is not assigned_to", "Return 403."],
        ["All tasks completed", "Mark onboarding as complete."],
    ], widths=[2.7, 4.6])

    h(doc, "Employee Offboarding - What HR Can Do", 2)
    bullets(doc, [
        "Start offboarding from POST /api/v1/employees/{id}/offboarding with employees:write permission.",
        "Select reason: resignation, termination, retirement, or contract_end.",
        "Set last working date.",
        "Record knowledge risk level for handover sensitivity.",
        "Record or later update exit interview notes.",
        "Track penalties_json for outstanding loans, notice-period penalties, or similar offboarding amounts.",
        "Review knowledge transfer status for high/critical knowledge-risk employees.",
        "Approve a knowledge-transfer bypass when handover cannot be completed and a valid reason is provided.",
        "Record a knowledge-transfer bypass penalty in penalties_json when tenant policy or manual HR decision requires it.",
        "Continue offboarding only when mandatory knowledge transfer is completed or bypassed with approval.",
        "Complete offboarding after exit actions are done.",
        "Deactivate the employee's user account as part of completion.",
        "Revoke active sessions and refresh tokens.",
        "Trigger agent revocation for the employee's devices.",
        "Create lifecycle event history for terminated or resigned status.",
    ])
    h(doc, "Employee Offboarding - System Sequence", 2)
    number(doc, [
        "EmployeeController.StartOffboarding receives employee id and StartOffboardingCommand.",
        "Authorization checks employees:write.",
        "OffboardingService loads the employee and validates active status.",
        "The service inserts offboarding_records with reason, last_working_date, knowledge_risk_level, and status initiated.",
        "The service creates the offboarding checklist, including manager-owned knowledge transfer and task reassignment.",
        "If knowledge_risk_level is high or critical, the knowledge transfer task is mandatory.",
        "If knowledge_risk_level is critical, the system triggers additional handover workflow.",
        "The service creates a lifecycle event with event_type terminated or resigned.",
        "EmployeeOffboardingStarted is published.",
        "Consumers notify HR/manager, forfeit or close leave balances, prepare final settlement, and revoke agent access where applicable.",
        "The service calculates penalties_json, including outstanding loans or notice period tracking where provided.",
        "If HR/Admin approves a knowledge-transfer bypass, the service requires a bypass reason.",
        "The service resolves the bypass penalty amount from manual input, tenant default policy, or zero for audit-only bypass.",
        "The service appends a knowledge_transfer_bypass item into penalties_json with amount, currency, reason, approver, timestamp, and source.",
        "Complete offboarding stores exit_interview_notes.",
        "Before completion, the service validates that high/critical knowledge transfer is either completed or bypassed with approval.",
        "Before completion, the service validates penalties_json is finalized for final-settlement review.",
        "Employee employment_status is changed to terminated and termination_date is set.",
        "User account is deactivated.",
        "All sessions and refresh tokens are revoked.",
        "Offboarding record status changes to completed.",
        "EmployeeTerminated event is published for downstream cleanup and finalization.",
    ])
    add_table(doc, ["Offboarding edge case", "System handling"], [
        ["Employee not active", "Offboarding start is rejected."],
        ["Critical knowledge risk", "Triggers additional handover workflow expectation."],
        ["Notice-period tracking", "Stored in penalties_json with expected end date."],
        ["Agent installed", "Agent Gateway revokes active agent registration/access."],
        ["Leave/payroll dependencies", "Leave closure and final settlement are downstream consumers; payroll itself remains Phase 2."],
    ], widths=[2.7, 4.6])
    h(doc, "Knowledge Transfer, Bypass, and Penalty Handling", 2)
    p(doc, "Knowledge transfer is treated as a manager-owned handover step inside offboarding. For high and critical risk employees, it is mandatory before offboarding completion unless an authorized HR/Admin user approves a bypass. Any bypass and related penalty are recorded in offboarding_records.penalties_json for audit and final-settlement review.")
    add_table(doc, ["Confirmed / added rule", "Detail"], [
        ["Knowledge risk levels", "offboarding_records.knowledge_risk_level supports low, medium, high, and critical."],
        ["Critical knowledge risk behaviour", "Critical employees trigger additional handover workflows."],
        ["Manager checklist responsibility", "The offboarding checklist includes Manager: knowledge transfer and task reassignment."],
        ["Mandatory handover", "High/critical risk offboarding cannot complete until knowledge transfer is completed or bypassed with approval."],
        ["Bypass approver", "HR/Admin or configured workflow approver can approve bypass."],
        ["Bypass reason", "A reason is required for every bypass."],
        ["Penalty storage", "offboarding_records.penalties_json stores outstanding loans, notice-period violations, asset recovery, and knowledge-transfer bypass penalties."],
        ["Penalty amount", "Amount can be manual, tenant-policy default, or zero for audit-only bypass."],
        ["Final pay interaction", "Phase 1 records final-settlement inputs; Phase 2 payroll can consume the finalized penalty items."],
    ], widths=[2.4, 4.9])
    add_table(doc, ["Knowledge risk level", "Required handling"], [
        ["low", "Knowledge transfer is optional unless the offboarding template requires it."],
        ["medium", "Knowledge transfer task is created and assigned to the manager."],
        ["high", "Knowledge transfer is mandatory before completion unless bypassed by HR/Admin approval."],
        ["critical", "Additional handover workflow is triggered; completion requires handover or approved bypass."],
    ], widths=[2.4, 4.9])

    h(doc, "Simple English Workflow Scope", 1)
    p(doc, "This section explains the scope in plain operational language: what the user can do, what the system does after that, and what happens in common edge cases.")
    for workflow in SIMPLE_WORKFLOWS:
        h(doc, workflow["title"], 2)
        h(doc, "What the user can do", 3)
        bullets(doc, workflow["user"])
        h(doc, "What the system does", 3)
        number(doc, workflow["system"])
        h(doc, "Common rules and edge cases", 3)
        add_table(doc, ["Case", "System handling"], workflow["edge"], widths=[2.5, 4.8])

    h(doc, "Permission Inventory", 1)
    p(doc, "The following permission codes were found in the repository documentation and code references. They are listed explicitly so the client can see the scope rather than only a summarized count.")
    perm_rows = []
    for idx, code in enumerate(PERMISSIONS, 1):
        perm_rows.append([idx, code])
    add_table(doc, ["#", "Permission code"], perm_rows, widths=[0.5, 6.8])

    h(doc, "Database Table Inventory", 1)
    p(doc, "The table inventory below is based on the schema catalog and module documentation. Phase 1 product tables total 133, with 5 additional Phase 1 Developer Platform tables.")
    rows = []
    for m in MODULES:
        rows.append([m["name"], m["phase"], len(m["tables"]), ", ".join(m["tables"])])
    add_table(doc, ["Module", "Phase", "Table count/list count", "Tables"], rows, widths=[1.5, 1.0, 1.0, 3.8])

    h(doc, "Detailed Module Appendix", 1)
    p(doc, "This appendix expands the scope from the module-level knowledge base. It includes each Phase 1 module's overview and end-to-end logic where available, so the client can inspect the operational details rather than only the summary.")
    for file_path in detail_files():
        add_markdown_file(doc, file_path, base_level=2)

    h(doc, "Out of Scope / Future Scope Boundaries", 1)
    add_table(doc, ["Area", "Boundary"], [
        ["Payroll", "Designed as Phase 2 module; not active Phase 1 build scope."],
        ["Performance", "Designed as Phase 2 module; not active Phase 1 build scope."],
        ["Full Skills and Learning", "Only Skills Core is Phase 1; courses, assessments, certifications, and development plans are Phase 2."],
        ["Documents", "Phase 2 module; file_records exist in Phase 1 as shared file registry."],
        ["Grievance", "Phase 2."],
        ["Expense", "Phase 2."],
        ["Reporting Engine", "Phase 2 table group, though Phase 1 analytics/reporting views exist in Productivity Analytics."],
        ["WMS development", "External system built by separate team; OneVo-HR provides bridge contracts and frontend consumption."],
        ["macOS agent", "Phase 2; Windows only in Phase 1."],
        ["Screen recording", "Phase 2/future; not Phase 1."],
        ["Silent capture", "Phase 2/future and subject to legal review; Phase 1 documented flow requires employee notification."],
        ["Customer developer API keys/webhooks", "Phase 2; distinct from the internal Developer Platform."],
    ], widths=[2.0, 5.3])

    h(doc, "Source Traceability", 1)
    bullets(doc, [
        "docs/onevo-phase1-scope.md",
        "docs/HR-Scope-Document-Phase1-Phase2.md",
        "backend/module-catalog.md",
        "database/schema-catalog.md",
        "modules/*/overview.md",
        "modules/agent-gateway/*",
        "modules/activity-monitoring/overview.md",
        "modules/identity-verification/overview.md",
        "developer-platform/*",
        "security/*",
        "frontend/*",
        "infrastructure/*",
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build()
